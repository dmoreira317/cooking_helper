## This is a program that prints in a document a desired recipe. You can search for one or get a random recipe.
import requests
import json
import pprint
import random
import sys
import html
import docx
import docx.shared
from docx.shared import Inches
from tkinter import ttk
from tkinter.filedialog import asksaveasfile

# maybe try infinite round of printing recipes? need to add a random name generator for each time as default. keep track? SOLVED WITH CHOOSE WHERE TO SAVE().
# join the save as window to my already generated doc, try to save the file first and return the file path to modify and save doc internally later. SOLVED WITH SAVE().
# make different documents? to be able to print more than one. SOLVED WITH SAVE().
# I need to add meal: null validation for unsuccesful searches. IMPORTANT one.
# save route and quit program validation
# quit_program = False # i need to complete this validation.

"""
#size of 'save as' window this is extra info for me, to create a window
root = Tk()
root.geometry("200x150")

to print a button on the window
# save button to save at any point
#btn = ttk.Button(root, text = "save", command = lambda : save())
#btn.pack(side = TOP, pady = 20)
"""

# functions definition
def save():
    files = [("Word files", ".docx")]
    file = asksaveasfile(filetypes = files, defaultextension = ".docx")
    print(file.name)
    return file.name

def print_doc(random_dict, ingredient_dict, img):
    user_print = input("\n--Would you like to print the recipe in a word document?\n// Please press enter to print or no to quit.\n")
    if user_print == "":
        # ask where to save
        save_route = save()
        
        # Printing word document
        recipe_doc = docx.Document()
        recipe_doc.add_paragraph("--------This is your meal of choice--------")
        recipe_doc.add_paragraph(html.unescape(random_dict["meals"][0]["strMeal"]))
        p = recipe_doc.add_paragraph()
        r = p.add_run()
        r.add_picture(img.name, width = Inches(1.25))
        recipe_doc.add_paragraph(html.unescape(random_dict["meals"][0]["strYoutube"]))
        recipe_doc.add_paragraph("\nList of Ingredients:")        
        for key in ingredient_dict:
            recipe_doc.add_paragraph(f"{key}, {ingredient_dict[key]}")
        recipe_doc.add_paragraph("Instructions:\n")
        recipe_doc.add_paragraph(html.unescape(random_dict["meals"][0]["strInstructions"]))
        recipe_doc.save(save_route)
        print("The document was created successfully!")
        print("Thanks for using this program, have a wonderful cooking session!")
    else:
        print("Thanks for using this program, have a wonderful cooking session!")

def get_img(random_dict):
    response = requests.get(random_dict["meals"][0]["strMealThumb"])
    img = open("sample_img.jpg", "wb") # opened with 'wb' to view it as bytes
    img.write(response.content)
    img.close()
    return img

def get_ingredients(random_dict):
    ingredient_dict = {}
    ingredient_up = "a"
    measurment_up = "a"
    ingredient = 1
    while ingredient_up != "" or measurment_up != "":
            try:
                ingredient_up = random_dict["meals"][0][f"strIngredient{ingredient}"]
                #print(ingredient_up)
                measurment_up = random_dict["meals"][0][f"strMeasure{ingredient}"]
                #print(measurment_up)
                ingredient_list_update = {ingredient_up: measurment_up}
                ingredient_dict.update(ingredient_list_update)
                ingredient += 1
            except:
                break
    
    #print(ingredient_dict)
    key_remove = ""
    if key_remove in ingredient_dict.keys():
        ingredient_dict.pop(key_remove)
    else:
        pass
    return ingredient_dict

def headers(random_dict):
    print("Meal: ", html.unescape(random_dict["meals"][0]["strMeal"]))
    print("Youtube video: ", html.unescape(random_dict["meals"][0]["strYoutube"]))
    print("Category: ", html.unescape(random_dict["meals"][0]["strCategory"]))

def list_of_ingredients_and_measurements(ingredient_dict):
    print("\n")
    print("List of ingredients:\n")
    for key in ingredient_dict:
        print(key, ": ", ingredient_dict[key])

def request_status_check(random_request):
    if random_request.status_code != 200:
        print("//There was a problem when loading the recipe, please re-start program.")
        quit_program = True
        recipe_quit = True
    else:
        quit_program = False

# First message to user
print("\n############################################################################################")
print("This is a cooking helper program. Choose a recipe and print it on a document. by @dlm317")
print("############################################################################################\n")

# Main program
def main():
    try:
        recipe_quit = False

        while recipe_quit == False:
            recipe_choice = input("Hello there, do you want to search for a recipe or get a random one?\nPress enter for a random recipe or type search to look for one.\n-->").lower()
            # Search for recipe
            if recipe_choice == "search":
                
                recipe_search = input("Please enter your desired recipe:").lower()
                search_request = requests.get(f"https://www.themealdb.com/api/json/v1/1/search.php?s={str(recipe_search)}")
                request_status_check(search_request)
                
                # If we get here, we can quit the recipe_quit loop 
                recipe_quit = True
                
                # Getting a recipe from the search function
                search_dict = json.loads(search_request.text)
                
                # Recipe headers
                headers(search_dict)

                # get recipe image from URL
                img = get_img(search_dict)
                            
                # Ingredients list
                ingredient_dict = get_ingredients(search_dict)
                            
                # On-screen ingredient list and measurments
                list_of_ingredients_and_measurements(ingredient_dict)
                
                # Recipe instructions
                print("\nInstructions:\n ", html.unescape(search_dict["meals"][0]["strInstructions"]))
                
                #print(random_request.text)
                
                # ask for user to print and save recipe
                print_doc(search_dict, ingredient_dict, img)
                                
            # Random request
            elif recipe_choice == "":
                
                random_request = requests.get("http://www.themealdb.com/api/json/v1/1/random.php")
                #print(random_request.text)
                
                #request status validation
                request_status_check(random_request)
                
                # If we get here, we can quit the recipe_quit loop
                recipe_quit = True
                        
                # Getting a recipe from the search function
                random_dict = json.loads(random_request.text)
                
                # Recipe headers
                headers(random_dict)
                
                # get recipe image from URL
                img = get_img(random_dict)
                            
                # Ingredients list
                ingredient_dict = get_ingredients(random_dict)
                
                # On-screen ingredient list and measurments
                list_of_ingredients_and_measurements(ingredient_dict)
                
                # Recipe instructions print
                print("\nInstructions:\n ", html.unescape(random_dict["meals"][0]["strInstructions"]))
                            
                # ask for user to print and save recipe
                print_doc(random_dict, ingredient_dict, img)

            # For invalid options
            else:
                print("Please select a valid option. Enter for random recipe or search to look for one.")
                continue

    except KeyboardInterrupt:
        print("Exiting program.")
        sys.exit(1)

if __name__ == "__main__":
    main()